import streamlit as st
import os
import docx2txt
import json
import requests
import docx
import re


# Fonction pour convertir DOCX en texte brut
def convert_docx_to_txt(input_path, output_path):
    text = docx2txt.process(input_path)
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(text)
    return output_path

# Fonction pour nettoyer les répétitions dans un fichier texte
def remove_repetitions(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    unique_lines = []
    seen_lines = set()
    for line in lines:
        if line.strip() not in seen_lines:
            unique_lines.append(line)
            seen_lines.add(line.strip())
    with open(output_file, 'w', encoding='utf-8') as file:
        file.writelines(unique_lines)
    return output_file

# Fonction pour interroger l'API
def ask_question(question, file_content=""):
    api_url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer wD0wgblQCZvPMzHoLj8cu6ZQOdQkibrR",
        "Content-Type": "application/json"
    }
    data = {
        "model": "mistral-medium-latest",
        "temperature": 0.7,
        "top_p": 1,
        "max_tokens": 1000,
        "messages": [{"role": "user", "content": f"{file_content}\n\n{question}"}]
    }
    reponse = requests.post(api_url, json=data, headers=headers)
    if reponse.status_code == 200:
        return reponse.json().get("choices")[0].get("message").get("content")
    return f"Erreur: {reponse.status_code}, {reponse.text}"

# Fonction pour nettoyer les réponses API et les formater en JSON
def parse_criteria(criteria_str):
    try:
        return json.loads(criteria_str)
    except json.JSONDecodeError:
        return criteria_str

# Fonction pour transformer un JSON imbriqué en dictionnaire
def parse_nested_json(json_str):
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        return json_str

# Fonction pour nettoyer et reformater la section "Tableau"
# Fonction pour nettoyer et reformater la section "Tableau"
def clean_tableau(tableau_str):
    try:
        # Essayer de convertir la chaîne mal formée en JSON
        tableau_dict = json.loads(tableau_str)
        # Si la conversion réussit, on retourne une liste de tableaux correctement formattés
        cleaned_tableau = []
        for key, value in tableau_dict.items():
            cleaned_tableau.append({
                "Tableau": key,
                "Titre": value.get("titre", "Titre inconnu"),
                "Description": value.get("Description", "Description indisponible")
            })
        return cleaned_tableau
    except json.JSONDecodeError:
        # Si la chaîne JSON est mal formée, renvoyer une valeur vide ou un message d'erreur
        return []

# Interface Streamlit
st.title("Écriture de CRF à partir de Protocole ")

# === Étape 1 : Conversion DOCX en texte brut ===
uploaded_file = st.file_uploader("Téléverser un document DOCX", type="docx")

if uploaded_file:
    # Vérifier si le dossier temp existe
    if not os.path.exists("temp"):
        os.makedirs("temp")

    # Sauvegarder le fichier DOCX uploadé
    input_path = os.path.join("temp", uploaded_file.name)
    with open(input_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Conversion du DOCX en texte brut
    txt_output_path = input_path.replace('.docx', '.txt')
    convert_docx_to_txt(input_path, txt_output_path)

    # Nettoyage des répétitions
    cleaned_txt_path = txt_output_path.replace('.txt', '_clean.txt')
    remove_repetitions(txt_output_path, cleaned_txt_path)

    # Téléchargement du fichier texte nettoyé
    with open(cleaned_txt_path, "rb") as file:
        st.download_button(
            label="Download Cleaned TXT",
            data=file,
            file_name=os.path.basename(cleaned_txt_path),
            mime="text/plain",
            key = "download_cleaned_txt"
        )

# === Étape 2 : Interrogation de l'API avec questions pré-écrites ===
if st.button("Récupérer les informations dans le Protocole"):
    # Chargement du fichier texte pour poser des questions
    with open(cleaned_txt_path, "r", encoding="utf-8") as file:
        file_content = file.read()

    questions_and_titles = [
        {
            "question": "Donne moi les criteres d'inclusion dans ce Protocole. Donne moi la reponse sous format json sans retour a la ligne comme ceci {'Critere 1': 'critere','Critere 2': 'critere','Critere 3': 'critere',...} En remplacent les guillemets simples par des guillemets double, Tu n'ajoutera aucun texte avant ou apres le code json, Do not add any other text outside the json code ",
            "title": "inclusion"},
        {
            "question": "Donne moi les criteres de non-inclusion dans ce Protocole? Donne moi la reponse sous format json sans retour a la ligne comme ceci {'Critere 1': 'critere','Critere 2': 'critere','Critere 3': 'critere',...} En remplacent les guillemets simples par des guillemets double, Tu n'ajoutera aucun texte avant ou apres le code json, Do not add any other text outside the json code  ",
            "title": "non_inclusion"},
        {
            "question": "Quelles sont les semaine des visite de suivis? Donne moi la reponse sous format json sans retour a la ligne comme ceci  {'Visite 1': 'Visite','Visite 2': 'Visite','Visite 3': 'Visite',...} En remplacent les guillemets simples par des guillemets double, Tu n'ajoutera aucun texte avant ou apres le code json, Do not add any other text outside the json code ",
            "title": "semaine"},
        {
            "question": "Donne moi L'acronyme de l'etude dans ce Protocole ? Donne moi la reponse sous format json sans retour a la ligne comme ceci {'acronyme': 'blabla'} En remplacent les guillemets simples par des guillemets double , Tu n'ajoutera aucun texte avant ou apres le code json, Do not add any other text outside the json code ",
            "title": "titre"},
        {
            "question": "Donne moi les tableaux de bilan biologique, hepatique, sanguins, etc.. necessaires à chaque visite ainsi qu'un résumé de ce qu'ils sont est sencé comporter dans ce Protocole? Donne moi la reponse sous format json sans retour a la ligne comme ceci {'Tableau 1','titre du tableaux','Description des valeurs névéssaire dans le tableau': '...'},{'Tableau 2','titre du tableaux','Description'': '...'}, En remplacent les guillemets simples par des guillemets double, Tu n'ajoutera aucun texte avant ou apres le code json, Do not add any other text outside the json code ",
            "title": "Tableau"}
    ]

    # Stocker les réponses API
    reponses_json = {}
    for item in questions_and_titles:
        reponse = ask_question(item["question"], file_content)
        reponses_json[item["title"]] = reponse
        st.write(f"reponse to '{item['question']}':\n{reponse}\n")

    # Sauvegarde des réponses API
    json_output_path = "reponses.json"
    with open(json_output_path, "w", encoding="utf-8") as output_file:
        json.dump(reponses_json, output_file, ensure_ascii=False, indent=4)

    # Téléchargement des réponses
    with open(json_output_path, "rb") as file:
        st.download_button(
            label="Download API reponses (JSON)",
            data=file,
            file_name="reponses.json",
            mime="application/json",
            key="download_api_responses"
        )

# === Étape 3 : Reformatage des réponses API et téléchargement ===
if st.button("Nettoyer les reponses"):
    # Reformatage des réponses API
    # Charger le fichier JSON initial
    with open("reponses.json", "r", encoding="utf-8") as file:
        json_initial = json.load(file)
        

    # Reformater les données JSON
    json_reformate = {
        "inclusion": json.loads(json_initial.get("inclusion", "{}")),
        "non_inclusion": json.loads(json_initial.get("non_inclusion", "{}")),
        "semaine": json.loads(json_initial.get("semaine", "{}")),
        "titre": json.loads(json_initial.get("titre", "{}")),
        "Tableau": clean_tableau(json_initial.get("Tableau", "{}"))  # Nettoyage spécifique pour "Tableau"
    }

    # Sauvegarde du fichier JSON reformatté
    json_clean_output_path = "reponses_clean.json"
    with open(json_clean_output_path, "w", encoding="utf-8") as f:
        json.dump(json_reformate, f, ensure_ascii=False, indent=4)


# === Étape 4 : Traitement des tableaux biologiques ===
if st.button("Générer les tableau nécessaire"):
    json_clean_output_path = "reponses_clean.json"
    # Charger les tableaux biologiques du JSON reformatté
    with open(json_clean_output_path, "r", encoding="utf-8") as file:
        json_data = json.load(file)

    tableaux = json_data.get("Tableau", [])
    nouveaux_tableaux = {}

    for index, tableau_info in enumerate(tableaux):
        titre = tableau_info.get("Titre", "Titre inconnu")
        description = tableau_info.get("Description", "Description indisponible")

        # Interroger l'API pour obtenir les détails des tableaux biologiques
        tableau_question = f"Sur la base du titre du tableau '{titre}' avec la description suivante : '{description}' et tes connaissances médicales, crée un tableau au format .json. Le tableau doit inclure les informations suivantes pour chaque ligne de test: - Le nom de la valeur à relever - l'unité de mesure habituelle pour cette valeur, Comme ceci : {{'Titre du tableau': 'titre','date': 'Date du bilan', 'results': [{{'test': 'Nom de la valeur à tester','value': 'Valeur obtenue','unité': 'unité de mesure normale'}}// Ajoutez d'autres tests ici]}}. En remplaçant les guillemets simples par des guillemets doubles. Tu n'ajoutera aucun texte avant ou apres le code json, Do not add any other text outside the json code "

        reponse = ask_question(tableau_question)
        nouveaux_tableaux[f"nouveau_tableau_{index + 1}"] = reponse

    # Sauvegarde des nouveaux tableaux
    nouveaux_tableaux_path = "nouveaux_tableaux.json"
    with open(nouveaux_tableaux_path, "w", encoding="utf-8") as output_file:
        json.dump(nouveaux_tableaux, output_file, ensure_ascii=False, indent=4)

    # Téléchargement des nouveaux tableaux
    with open(nouveaux_tableaux_path, "rb") as file:
        st.download_button(
            label="Download Generated Tables (JSON)",
            data=file,
            file_name="nouveaux_tableaux.json",
            mime="application/json",
            key="download_generated_tables"
        )

# === Étape 5 : Nettoyage des tableaux biologiques ===
if st.button("Nettoyer les tableau"):
    # Charger les nouveaux tableaux
    nouveaux_tableaux_path = "nouveaux_tableaux.json"
    with open(nouveaux_tableaux_path, "r", encoding="utf-8") as file:
        nouveaux_tableaux = json.load(file)

    tableaux_propres = {}
    for key, value in nouveaux_tableaux.items():
        tableau_propre = parse_nested_json(value)
        tableaux_propres[key] = tableau_propre

        # Sauvegarde des tableaux propres
        tableaux_propres_path = "tableaux_propres.json"
        with open(tableaux_propres_path, "w", encoding="utf-8") as output_file:
            json.dump(tableaux_propres, output_file, ensure_ascii=False, indent=4)



# === Étape 6 : Remplissage du CRF avec les données JSON ===
if st.button("Générer le CRF"):
    import json
    from docx import Document
    from copy import deepcopy

    # Fonction pour modifier le document CRF avec les données JSON
    def modify_crf_document():
        # Charger le document Word existant
        doc = Document('CRF_0.docx')

        # Charger les données depuis le fichier JSON
        with open('reponses_clean.json', 'r', encoding='utf-8') as file:
            data = json.load(file)

        # Modification des critères d'inclusion
        table_inclusion = doc.tables[3]
        inclusion_criteres = data['inclusion']
        num_criteres_inclusion = len(inclusion_criteres)

        lignes_intermediaires = len(table_inclusion.rows) - 2  # Exclure la première et dernière ligne

        # Ajustement des lignes selon le nombre de critères d'inclusion
        if num_criteres_inclusion < lignes_intermediaires:
            for _ in range(lignes_intermediaires - num_criteres_inclusion):
                table_inclusion._element.remove(table_inclusion.rows[-2]._element)
        elif num_criteres_inclusion > lignes_intermediaires:
            for _ in range(num_criteres_inclusion - lignes_intermediaires):
                new_row = deepcopy(table_inclusion.rows[1]._element)
                table_inclusion._element.insert(-2, new_row)

        # Itérer à travers les critères d'inclusion en utilisant les clés du dictionnaire
        for i, (key, value) in enumerate(inclusion_criteres.items()):
            table_inclusion.cell(i + 1, 0).text = value  # Remplir avec la description du critère

        # Modification des critères de non-inclusion
        table_n_inclusion = doc.tables[4]
        non_inclusion_criteres = data['non_inclusion']
        num_criteres_n_inclusion = len(non_inclusion_criteres)
        lignes_intermediaires = len(table_n_inclusion.rows) - 2

        if num_criteres_n_inclusion < lignes_intermediaires:
            for _ in range(lignes_intermediaires - num_criteres_n_inclusion):
                table_n_inclusion._element.remove(table_n_inclusion.rows[-2]._element)
        elif num_criteres_n_inclusion > lignes_intermediaires:
            for _ in range(num_criteres_n_inclusion - lignes_intermediaires):
                new_row = deepcopy(table_n_inclusion.rows[1]._element)
                table_n_inclusion._element.insert(-2, new_row)

        # Itérer à travers les critères de non-inclusion
        for i, (key, value) in enumerate(non_inclusion_criteres.items()):
            table_n_inclusion.cell(i + 1, 0).text = value  # Remplir avec la description du critère

        # Ajout des visites
        nombre_visites = len(data['semaine'])
        table_15 = doc.tables[15]
        table_16 = doc.tables[16]
        paragraph_46 = doc.paragraphs[44]

        for visite_num in range(2, nombre_visites + 1):
            new_paragraph = paragraph_46.insert_paragraph_before()
            new_table_16 = deepcopy(table_16._element)
            new_paragraph._element.addnext(new_table_16)
            new_paragraph = new_paragraph.insert_paragraph_before()
            new_paragraph.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            new_table_15 = deepcopy(table_15._element)
            new_paragraph._element.addnext(new_table_15)
            new_paragraph = new_paragraph.insert_paragraph_before()
            new_paragraph.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            new_paragraph.add_run(f"VISITE {visite_num}").bold = True
            new_paragraph.style = 'TITRE PAGE'

        # Ajout des tableaux biologiques
        with open('tableaux_propres.json', 'r', encoding='utf-8') as file:
            data_tableau = json.load(file)

        table_model = doc.tables[-1]

        def modify_and_duplicate_table(table_model, table_data_tableau):
            for i, result in enumerate(table_data_tableau['results']):
                if len(table_model.rows) <= i + 4:
                    new_row = deepcopy(table_model.rows[3]._element)
                    table_model._element.append(new_row)

            table_model.cell(0, 0).text = f"Bilan {table_data_tableau['Titre du tableau']}"
            table_model.cell(1, 0).text = f"Bilan {table_data_tableau['Titre du tableau']}"
            table_model.cell(2, 1).text = f"Date du Bilan {table_data_tableau['Titre du tableau']}"

            for i, result in enumerate(table_data_tableau['results']):
                table_model.cell(i + 3, 1).text = result['test']
                table_model.cell(i + 3, 2).text = "|__|__|, |__|__|", result['unité']

        for key, table_info in data_tableau.items():
            doc.add_paragraph()
            new_table = deepcopy(table_model._element)
            doc._element.append(new_table)
            modify_and_duplicate_table(doc.tables[-1], table_info)

        # Sauvegarder le document modifié
        doc.save("CRF_Modif.docx")
        return "CRF_Modif.docx"

    # Modifier et générer le CRF modifié
    crf_file = modify_crf_document()

    # Télécharger le CRF modifié via Streamlit
    with open(crf_file, "rb") as file:
        st.download_button(
            label="Download Modified CRF",
            data=file,
            file_name="CRF_Modif.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_modified_crf"
        )
